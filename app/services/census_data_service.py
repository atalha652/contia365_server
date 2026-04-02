"""
Census Data Extraction Service for Contia365
Regex-based parser for Spanish Certificado de Situación Censal.
No AI or external API required.
"""

import io
import re
from typing import Optional, List

import pdfplumber


# ─────────────────────────── helpers ────────────────────────────────────────

def _get(text: str, *patterns: str) -> Optional[str]:
    for pattern in patterns:
        m = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        if m:
            return m.group(1).strip()
    return None


def _get_float(text: str, *patterns: str) -> Optional[float]:
    raw = _get(text, *patterns)
    if raw is None:
        return None
    raw = re.sub(r"[€$£\s]", "", raw).replace(",", "")
    try:
        return float(raw)
    except ValueError:
        return None


def _normalise_date(d: Optional[str]) -> Optional[str]:
    """dd/mm/yyyy or dd-mm-yyyy → yyyy-mm-dd"""
    if not d:
        return None
    d = d.strip()
    m = re.match(r"(\d{2})[/\-](\d{2})[/\-](\d{4})", d)
    if m:
        return f"{m.group(3)}-{m.group(2)}-{m.group(1)}"
    # "26 de marzo de 2026" style
    months = {
        "enero": "01", "febrero": "02", "marzo": "03", "abril": "04",
        "mayo": "05", "junio": "06", "julio": "07", "agosto": "08",
        "septiembre": "09", "octubre": "10", "noviembre": "11", "diciembre": "12",
    }
    m2 = re.match(r"(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})", d, re.IGNORECASE)
    if m2:
        month = months.get(m2.group(2).lower())
        if month:
            return f"{m2.group(3)}-{month}-{m2.group(1).zfill(2)}"
    if re.match(r"\d{4}-\d{2}-\d{2}", d):
        return d
    return None


def _detect_document_type(text: str) -> str:
    t = text.lower()
    if "situación censal" in t or "situacion censal" in t or "certificado censal" in t:
        return "Certificado de Situación Censal"
    if "modelo 100" in t or ("irpf" in t and "cuota" in t):
        return "Modelo 100 (IRPF)"
    if "census tax declaration" in t or "household" in t:
        return "Census Tax Declaration"
    return "Unknown"


# ─────────────────────────── text extraction ────────────────────────────────

def _extract_text_from_pdf(file_bytes: bytes) -> str:
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
            for table in (page.extract_tables() or []):
                for row in table:
                    if not row:
                        continue
                    cells = [str(c).strip() if c else "" for c in row]
                    if any(cells):
                        parts.append("TABLE_ROW|" + "|".join(cells))
    return "\n".join(parts)


def _extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Run: pip install python-docx")
    doc = Document(io.BytesIO(file_bytes))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = "TABLE_ROW|" + "|".join(c.text.strip() for c in row.cells)
            if row_text.strip("TABLE_ROW|"):
                lines.append(row_text)
    return "\n".join(lines)


def extract_text_from_file(file_bytes: bytes, content_type: str, filename: str) -> str:
    fname = filename.lower()
    if content_type == "application/pdf" or fname.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes)
    if content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ) or fname.endswith((".docx", ".doc")):
        return _extract_text_from_docx(file_bytes)
    raise ValueError(f"Unsupported file type: {content_type or filename}")


# ─────────────────────────── section parsers ────────────────────────────────

def _parse_taxpayer(text: str) -> dict:
    # NIF — look for explicit label first, then bare DNI/NIE pattern
    nif = _get(text,
               r"NIF/NIE[:\s]+([A-Z0-9]{8,9})",
               r"\bNIF[:\s]+([A-Z0-9]{8,9})",
               r"\bNIE[:\s]+([A-Z0-9]{8,9})")
    if not nif:
        m = re.search(r"\b([0-9]{8}[A-Z])\b", text)
        if m:
            nif = m.group(1)
    if not nif:
        m = re.search(r"\b([XYZ][0-9]{7}[A-Z])\b", text)
        if m:
            nif = m.group(1)

    # Full name — "Nombre o Razón Social: NAME" or "NOMBRE/RAZON SOCIAL: NAME"
    name = _get(text,
                r"NOMBRE/RAZON\s+SOCIAL[:\s]+(.+)",
                r"Nombre\s+o\s+Raz[oó]n\s+Social[:\s]+(.+)")

    # Fiscal address — grab lines after "Domicilio fiscal en España" up to "Residente"
    addr_block_match = re.search(
        r"Domicilio\s+fiscal\s+en\s+Espa[ñn]a\s*\n(.*?)(?=\nResidente|\nSITUACI|\nIDENTIF|\Z)",
        text, re.IGNORECASE | re.DOTALL
    )
    raw_addr = None
    city = None
    postal = None
    province = None

    if addr_block_match:
        addr_block = addr_block_match.group(1)
        addr_lines = [l.strip() for l in addr_block.splitlines() if l.strip()]

        # The address proper ends before "Localidad/Población" token
        addr_parts = []
        for line in addr_lines:
            if re.match(r"Localidad/Poblaci[oó]n", line, re.I):
                # city + postal + province are on this line
                # "Localidad/Población PALMA DE MALLORCA 07002 PALMA (ILLES BALEARS)"
                # or split: "MALLORCA 07002 PALMA (ILLES BALEARS)"
                loc_text = re.sub(r"Localidad/Poblaci[oó]n\s*", "", line, flags=re.I).strip()
                # extract postal
                pm = re.search(r"\b(\d{5})\b", loc_text)
                if pm:
                    postal = pm.group(1)
                    # city = everything before the postal
                    city = loc_text[:pm.start()].strip()
                    # province = inside parens after postal
                    prov_m = re.search(r"\d{5}\s+\w+\s+\(([^)]+)\)", loc_text)
                    if prov_m:
                        province = prov_m.group(1).strip()
                break
            else:
                addr_parts.append(line)

        # If city was split across lines (e.g. "PALMA DE\nMALLORCA 07002...")
        # the addr_parts may contain the city prefix — detect and move it
        if addr_parts:
            last = addr_parts[-1]
            # if last addr line has a postal in it, it's actually city+postal
            pm2 = re.search(r"\b(\d{5})\b", last)
            if pm2 and not postal:
                postal = pm2.group(1)
                city_prefix = last[:pm2.start()].strip()
                addr_parts = addr_parts[:-1]
                prov_m2 = re.search(r"\d{5}\s+\w+\s+\(([^)]+)\)", last)
                if prov_m2:
                    province = prov_m2.group(1).strip()
                if city_prefix:
                    city = city_prefix
            raw_addr = " ".join(addr_parts).strip() or None

        # city may still have a newline artifact — clean it
        if city:
            city = re.sub(r"\s+", " ", city).strip()

    # fallback postal from document if still missing
    if not postal:
        all_postals = re.findall(r"\b(\d{5})\b", text)
        if len(all_postals) >= 2:
            postal = all_postals[1]

    # Resident status
    resident_status = bool(re.search(r"Residente[:\s]+SI", text, re.IGNORECASE))

    return {
        "nif_nie": nif,
        "full_name": name,
        "fiscal_address": {
            "address_line": raw_addr,
            "postal_code": postal,
            "city": city,
            "province": province,
        },
        "resident_status": True if resident_status else None,
    }


def _parse_professional_registration(text: str) -> Optional[dict]:
    # VAT regime: "- General  08-05-2018" under "Regímenes aplicables"
    vat_regime = _get(text,
                      r"R[eé]gimenes?\s+aplicables[^\n]*\n\s*-\s+(\w+)",
                      r"R[eé]gimen\s*(?:de\s*)?IVA[:\s]+(.+)")

    # IRPF method: "- Estimación directa simplificada desde: 08-05-2018"
    irpf_method = _get(text,
                       r"-\s+(Estimaci[oó]n\s+directa\s+\w+)\s+desde",
                       r"M[eé]todo\s*IRPF[:\s]+(.+)")

    activities = _parse_economic_activities(text)

    if not vat_regime and not irpf_method and not activities:
        return None

    return {
        "vat_regime": vat_regime,
        "irpf_method": irpf_method,
        "economic_activities": activities,
    }


def _parse_economic_activities(text: str) -> List[dict]:
    """
    Parse the ACTIVIDADES ECONÓMICAS table.

    pdfplumber can produce TABLE_ROW lines in two formats:

    Format A — full row per activity (ideal):
        TABLE_ROW|Empresarial|967.2 - Escuelas...|1|08/05/2018|A03

    Format B — description split across rows (common with wrapped cells):
        TABLE_ROW|Empresarial|967.2 - Escuelas y Servicios de|1||
        TABLE_ROW||Perfeccionamiento del Deporte||08/05/2018|A03

    Plain-text fallback:
        Empresarial
        967.2 - Escuelas y Servicios de
        Perfeccionamiento del Deporte 1 08/05/2018 A03
        Profesional
        841 - Naturópata, Acupuntores y Otros
        Profesionales Parasanitarios 1 08/05/2018 A05
    """
    activities = []

    # ── Strategy 1: collect ALL TABLE_ROW lines in the activities block ──
    # Gather them first, then merge split rows
    table_lines = []
    in_activities = False
    for line in text.splitlines():
        if re.search(r"ACTIVIDADES\s+ECON[OÓ]MICAS", line, re.I):
            in_activities = True
            continue
        if in_activities and re.search(r"OBLIGACIONES|Delegaci[oó]n|App\s+AEAT|Documento\s+firmado", line, re.I):
            in_activities = False
        if in_activities and line.startswith("TABLE_ROW|"):
            parts = [p.strip() for p in line.split("|")[1:]]
            # skip pure header rows
            if re.search(r"secci[oó]n|grupo|epígrafe|n[oº]\s*actividades|fecha\s*alta|c[oó]digo\s*de",
                         "|".join(parts), re.I):
                continue
            table_lines.append(parts)

    if table_lines:
        # Merge rows: a row "belongs" to the previous if it has an activity code
        # but the previous row was missing one
        merged = []
        for parts in table_lines:
            # Find activity code in any cell
            act_code = next((p for p in parts if re.match(r"^A\d+$", p)), None)
            date_val = next((p for p in parts if re.match(r"^\d{2}/\d{2}/\d{4}$", p)), None)

            if act_code and date_val:
                # This is a complete or terminal row — find section and description
                section = next(
                    (p for p in parts if p and not re.match(r"^A\d+$|^\d+$|^\d{2}/\d{2}/\d{4}$", p)
                     and re.match(r"^(Empresarial|Profesional|Agr[ií]cola|Art[ií]stica)", p, re.I)),
                    None
                )
                # description = largest non-section, non-code, non-date, non-count cell
                desc_candidates = [
                    p for p in parts
                    if p
                    and not re.match(r"^A\d+$", p)
                    and not re.match(r"^\d{2}/\d{2}/\d{4}$", p)
                    and not re.match(r"^(Empresarial|Profesional|Agr[ií]cola|Art[ií]stica)$", p, re.I)
                    and not re.match(r"^\d$", p)
                ]
                raw_desc = " ".join(desc_candidates).strip()

                # If previous merged entry is missing its code, append description to it
                if merged and merged[-1]["activity_type_code"] is None:
                    prev = merged[-1]
                    prev["description"] = (prev["description"] + " " + raw_desc).strip()
                    prev["start_date"] = _normalise_date(date_val)
                    prev["activity_type_code"] = act_code
                    if not prev["section"] and section:
                        prev["section"] = section
                else:
                    merged.append({
                        "section": section,
                        "raw_desc": raw_desc,
                        "start_date": _normalise_date(date_val),
                        "activity_type_code": act_code,
                    })
            else:
                # Incomplete row — start a pending entry
                section = next(
                    (p for p in parts if p and re.match(
                        r"^(Empresarial|Profesional|Agr[ií]cola|Art[ií]stica)", p, re.I)),
                    None
                )
                desc_candidates = [
                    p for p in parts
                    if p
                    and not re.match(r"^(Empresarial|Profesional|Agr[ií]cola|Art[ií]stica)$", p, re.I)
                    and not re.match(r"^\d$", p)
                ]
                raw_desc = " ".join(desc_candidates).strip()
                if raw_desc or section:
                    merged.append({
                        "section": section,
                        "raw_desc": raw_desc,
                        "start_date": None,
                        "activity_type_code": None,
                    })

        # Convert merged entries to final activity dicts
        for entry in merged:
            if not entry.get("activity_type_code"):
                continue  # incomplete, skip
            raw_desc = entry.get("raw_desc") or entry.get("description") or ""
            code_match = re.match(r"^([\d\.]+)\s*[-–]\s*(.+)$", raw_desc)
            if code_match:
                code = code_match.group(1)
                description = code_match.group(2).strip()
            else:
                code = None
                description = raw_desc
            activities.append({
                "section": entry.get("section"),
                "code": code,
                "description": description,
                "start_date": entry.get("start_date"),
                "activity_type_code": entry["activity_type_code"],
            })

        if activities:
            return activities

    # ── Strategy 2: plain-text block ──
    block_match = re.search(
        r"ACTIVIDADES\s+ECON[OÓ]MICAS\s*\n(.*?)(?=OBLIGACIONES|Delegaci[oó]n|App\s+AEAT|Documento\s+firmado|\Z)",
        text, re.IGNORECASE | re.DOTALL
    )
    if not block_match:
        return activities

    block = block_match.group(1)
    section = None
    pending_desc = []

    for line in block.splitlines():
        line = line.strip()
        if not line or line.startswith("TABLE_ROW"):
            continue

        # Section header on its own line
        if re.match(r"^(Empresarial|Profesional|Agr[ií]cola|Art[ií]stica)$", line, re.I):
            section = line
            pending_desc = []
            continue

        # Line starting with section name followed by IAE code inline
        # e.g. "Empresarial 967.2 - Escuelas..."
        inline = re.match(
            r"^(Empresarial|Profesional|Agr[ií]cola|Art[ií]stica)\s+([\d\.]+\s*[-–].+)$",
            line, re.I
        )
        if inline:
            section = inline.group(1)
            pending_desc = [inline.group(2).strip()]
            continue

        # Skip column headers
        if re.search(r"Secci[oó]n|Grupo/Ep[ií]grafe|N[oº]\s*actividades|Fecha\s*alta|C[oó]digo\s*de", line, re.I):
            continue

        # Terminal line: ends with COUNT  dd/mm/yyyy  A0x
        m = re.match(r"^(.*?)\s+\d\s+(\d{2}/\d{2}/\d{4})\s+(A\d+)\s*$", line)
        if m:
            desc_part = m.group(1).strip()
            full_desc = " ".join(pending_desc + ([desc_part] if desc_part else [])).strip()
            pending_desc = []

            code_match = re.match(r"^([\d\.]+)\s*[-–]\s*(.+)$", full_desc)
            activities.append({
                "section": section,
                "code": code_match.group(1) if code_match else None,
                "description": code_match.group(2).strip() if code_match else full_desc,
                "start_date": _normalise_date(m.group(2)),
                "activity_type_code": m.group(3),
            })
            continue

        # Continuation line
        if not re.search(r"Delegaci[oó]n|App\s+AEAT|Documento", line, re.I):
            pending_desc.append(line)

    return activities


def _parse_periodic_obligations(text: str) -> List[dict]:
    """
    Parse OBLIGACIONES PERIÓDICAS section.

    AEAT format (no modelo numbers in plain text, just descriptions):
        MODELO                          PERIODICIDAD
        IRPF-ISS. RET.ARREND...         TRIMESTRAL
        IRPF PAGO FRACCIONADO...        TRIMESTRAL
        IMPUESTO SOBRE EL VALOR...      TRIMESTRAL

    We map known descriptions to their modelo numbers.
    """
    DESCRIPTION_TO_MODELO = {
        r"RET\.?ARREND\.?INMUEBLES\s+URBANOS": "115",
        r"IRPF.*PAGO\s+FRACCIONADO": "130",
        r"IMPUESTO\s+SOBRE\s+EL\s+VALOR\s+A[ÑN]ADIDO": "303",
        r"RETENCIONES.*TRABAJO\s+PERSONAL": "111",
        r"DECLARACI[OÓ]N\s+ANUAL.*IVA": "390",
        r"DECLARACI[OÓ]N\s+ANUAL.*IRPF": "190",
    }

    obligations = []

    # Find the OBLIGACIONES PERIÓDICAS block
    block_match = re.search(
        r"OBLIGACIONES\s+PERI[OÓ]DICAS\s*\n(.*?)(?=\nY\s+para\s+que|\nDocumento\s+firmado|\Z)",
        text, re.IGNORECASE | re.DOTALL
    )
    if not block_match:
        return obligations

    block = block_match.group(1)

    for line in block.splitlines():
        line = line.strip()
        if not line:
            continue
        # Skip header line
        if re.match(r"^MODELO\s+PERIODICIDAD$", line, re.I):
            continue

        # Each line: "DESCRIPTION   PERIODICITY"
        # Periodicity is always at the end: TRIMESTRAL / MENSUAL / ANUAL
        period_match = re.search(r"\b(TRIMESTRAL|MENSUAL|ANUAL)\s*$", line, re.I)
        if not period_match:
            continue

        periodicity = period_match.group(1).upper()
        description = line[:period_match.start()].strip()

        # Resolve modelo number
        modelo = None
        for pattern, num in DESCRIPTION_TO_MODELO.items():
            if re.search(pattern, description, re.I):
                modelo = num
                break

        obligations.append({
            "modelo": modelo,
            "description": description,
            "periodicity": periodicity,
        })

    return obligations


def _parse_income_and_expenses(text: str) -> Optional[dict]:
    revenue = _get_float(text,
                         r"(?:Total\s*)?(?:Ingresos|Revenue)[:\s€]+([\d,\.]+)",
                         r"Total\s*Revenue[:\s€]+([\d,\.]+)")
    expenses = _get_float(text,
                          r"(?:Total\s*)?(?:Gastos\s*Deducibles|Deductible\s*Expenses)[:\s€]+([\d,\.]+)")
    net = _get_float(text,
                     r"(?:Beneficio\s*Neto|Net\s*Profit)[:\s€]+([\d,\.]+)")
    withholdings = _get_float(text,
                              r"(?:Retenciones\s*Acumuladas|Accumulated\s*Withholdings)[:\s€]+([\d,\.]+)")
    if all(v is None for v in [revenue, expenses, net, withholdings]):
        return None
    return {
        "total_revenue_period": revenue,
        "total_deductible_expenses": expenses,
        "net_profit": net,
        "accumulated_withholdings_received": withholdings,
    }


def _parse_tax_calculation(text: str) -> Optional[dict]:
    taxable_base = _get_float(text,
                              r"Base\s*Imponible[^:]*:[:\s€]+([\d,\.]+)",
                              r"Taxable\s*Base[^:]*:[:\s€]+([\d,\.]+)")
    tax_quota = _get_float(text,
                           r"Cuota\s*[ÍI]ntegra[^:]*:[:\s€]+([\d,\.]+)",
                           r"Tax\s*Quota[^:]*:[:\s€]+([\d,\.]+)")
    final_tax = _get_float(text,
                           r"Cuota\s*L[íi]quida[^:]*:[:\s€]+([\d,\.]+)",
                           r"Final\s*Tax[^:]*:[:\s€]+([\d,\.]+)")
    withholdings_paid = _get_float(text,
                                   r"Retenciones?\s*Pagadas?[:\s€]+([\d,\.]+)",
                                   r"Withholdings?\s*Paid[:\s€]+([\d,\.]+)")
    result_amount = _get_float(text,
                               r"Resultado[:\s]+(?:A\s*pagar|A\s*devolver)[:\s€]*([\d,\.]+)",
                               r"Result[:\s]+(?:Refund|Payment)[:\s€]*([\d,\.]+)")
    result_raw = _get(text,
                      r"Resultado[:\s]+(A\s*pagar|A\s*devolver)",
                      r"Result[:\s]+(Refund|Payment)")
    result_type = None
    if result_raw:
        result_type = "Refund" if re.search(r"refund|devolver", result_raw, re.I) else "Payment"

    if all(v is None for v in [taxable_base, tax_quota, final_tax]):
        return None
    return {
        "taxable_base": taxable_base,
        "tax_quota": tax_quota,
        "final_tax": final_tax,
        "withholdings_paid": withholdings_paid,
        "result_amount": result_amount,
        "result_type": result_type,
    }


def _parse_household_data(text: str) -> Optional[dict]:
    members = []
    for line in text.splitlines():
        if not line.startswith("TABLE_ROW|"):
            continue
        parts = [p.strip() for p in line.split("|")[1:] if p.strip()]
        if re.search(r"member\s*name|age|occupation|income|secci[oó]n|n[oº]\s*actividades", 
                     "|".join(parts), re.I):
            continue
        if len(parts) >= 3 and re.match(r"[A-Za-záéíóúÁÉÍÓÚñÑ]", parts[0]):
            try:
                members.append({
                    "name": parts[0],
                    "age": int(parts[1]) if parts[1].isdigit() else None,
                    "occupation": parts[2] if len(parts) > 2 else None,
                    "annual_income": float(parts[3].replace(",", "")) if len(parts) > 3 else None,
                })
            except (ValueError, AttributeError):
                pass

    total = _get_float(text,
                       r"Total\s*Household\s*Income[:\s€]+([\d,\.]+)",
                       r"Ingresos\s*Totales\s*(?:del\s*Hogar)?[:\s€]+([\d,\.]+)")

    if not members and total is None:
        return None
    return {"members": members, "total_household_income": total}


# ─────────────────────────── main entry point ───────────────────────────────

def parse_census_data_from_text(raw_text: str) -> dict:
    doc_type = _detect_document_type(raw_text)

    # Issue date: "26 de marzo de 2026" or "dd/mm/yyyy"
    issue_date = _normalise_date(
        _get(raw_text,
             r"con\s+fecha\s+(\d{1,2}\s+de\s+\w+\s+de\s+\d{4})",
             r"Fecha\s*de\s*Emisi[oó]n[:\s]+([\d\-/]+)",
             r"Issue\s*Date[:\s]+([\d\-/]+)")
    )

    # CSV code: "8NMSUWVQQNEJ83LD" — 16 uppercase alphanumeric chars
    csv_code = _get(raw_text,
                    r"C[oó]digo\s+Seguro\s+Verificaci[oó]n\s*\n?\s*([A-Z0-9]{16})",
                    r"\bCSV[:\s]+([A-Z0-9]{10,})")

    aeat_ref = _get(raw_text,
                    r"N[Oº]\s*DE\s*REFERENCIA[:\s]+(\S+)",
                    r"Referencia[:\s]+(\S+)",
                    r"Reference[:\s]+(\S+)")

    return {
        "document_metadata": {
            "document_type": doc_type,
            "official_name": _get(raw_text,
                                  r"^(CERTIFICADO\s+DE\s+SITUACI[OÓ]N\s+CENSAL)",
                                  r"^(Agencia Tributaria[^\n]+)") or doc_type,
            "issue_date": issue_date,
            "csv_code": csv_code,
            "aeat_reference": aeat_ref,
        },
        "taxpayer_identity": _parse_taxpayer(raw_text),
        "professional_registration": _parse_professional_registration(raw_text),
        "periodic_tax_obligations": _parse_periodic_obligations(raw_text),
        "income_and_expenses_summary": _parse_income_and_expenses(raw_text),
        "tax_calculation": _parse_tax_calculation(raw_text),
        "household_data": _parse_household_data(raw_text),
        "platform_verification": {
            "verification_status": "PENDING",
            "verified_at": None,
            "needs_renewal_at": None,
        },
    }


def build_ocr_confidence(raw_text: str) -> float:
    key_terms = ["NIF", "IRPF", "IVA", "ALTA", "Cuota", "Deducci", "Imponible",
                 "Tributaria", "Censal", "Retenci", "IAE", "Modelo", "AEAT"]
    hits = sum(1 for t in key_terms if t.lower() in raw_text.lower())
    return round(min(hits / len(key_terms), 1.0), 2)
